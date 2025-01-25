
from pyexpat.errors import messages
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Table, TableStyle,Frame
from io import BytesIO
import os
import pandas as pd
from django.conf import settings
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from .forms import PDFForm,quotationform,pricescheduleform
from .models import *
from django.shortcuts import get_object_or_404, render,redirect
from reportlab.lib.enums import TA_JUSTIFY,TA_CENTER
from reportlab.platypus import Spacer
from reportlab.platypus import Image

from django.core.files.base import ContentFile
from PyPDF2 import PdfMerger

from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import pandas as pd
import os
from docx.shared import RGBColor

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth


def rupee_format_get_d(x_str):
    arr_1 = ["One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", ""]
    return arr_1[int(x_str) - 1] if int(x_str) > 0 else ""

def rupee_format_get_t(x_str):
    arr1 = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
    arr2 = ["", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    result = ""

    if int(x_str[0]) == 1:
        result = arr1[int(x_str[1])]
    else:
        if int(x_str[0]) > 0:
            result = arr2[int(x_str[0]) - 1]
        result += rupee_format_get_d(x_str[1])

    return result

def rupee_format_get_h(x_str_h, x_lp):
    x_str_h = x_str_h.zfill(3)  # Pad to ensure it's 3 digits
    result = ""

    if int(x_str_h) < 1:
        return ""

    if x_str_h[0] != "0":
        if x_lp > 0:
            result = rupee_format_get_d(x_str_h[0]) + " Lakh "
        else:
            result = rupee_format_get_d(x_str_h[0]) + " Hundred "

    if x_str_h[1] != "0":
        result += rupee_format_get_t(x_str_h[1:])
    else:
        result += rupee_format_get_d(x_str_h[2])

    return result

def rupee_format(s_num):
    arr_place = ["", "", " Thousand ", " Lakhs ", " Crores ", " Trillion ", "", "", "", ""]
    if s_num == "":
        return ""

    x_num_str = str(s_num).strip()

    if x_num_str == "":
        return ""

    if float(x_num_str) > 999999999.99:
        return "Digit exceeds Maximum limit"

    x_dp_int = x_num_str.find(".")

    x_rstr_paisas = ""  # Initialize x_rstr_paisas here

    if x_dp_int > 0:
        if len(x_num_str) - x_dp_int == 1:
            x_rstr_paisas = rupee_format_get_t(x_num_str[x_dp_int + 1:x_dp_int + 2] + "0")
        elif len(x_num_str) - x_dp_int > 1:
            x_rstr_paisas = rupee_format_get_t(x_num_str[x_dp_int + 1:x_dp_int + 3])

        x_num_str = x_num_str[:x_dp_int]

    x_f = 1
    x_rstr = ""
    x_lp = 0

    while x_num_str != "":
        if x_f >= 2:
            x_temp = x_num_str[-2:]
        else:
            if len(x_num_str) == 2:
                x_temp = x_num_str[-2:]
            elif len(x_num_str) == 1:
                x_temp = x_num_str[-1:]
            else:
                x_temp = x_num_str[-3:]

        x_str_temp = ""

        if int(x_temp) > 99:
            x_str_temp = rupee_format_get_h(x_temp, x_lp)
            if "Lac" not in x_str_temp:
                x_lp += 1
        elif int(x_temp) <= 99 and int(x_temp) > 9:
            x_str_temp = rupee_format_get_t(x_temp)
        elif int(x_temp) < 10:
            x_str_temp = rupee_format_get_d(x_temp)

        if x_str_temp != "":
            x_rstr = x_str_temp + arr_place[x_f] + x_rstr

        if x_f == 2:
            if len(x_num_str) == 1:
                x_num_str = ""
            else:
                x_num_str = x_num_str[:-2]
        elif x_f == 3:
            if len(x_num_str) >= 3:
                x_num_str = x_num_str[:-2]
            else:
                x_num_str = ""
        elif x_f == 4:
            x_num_str = ""
        else:
            if len(x_num_str) <= 2:
                x_num_str = ""
            else:
                x_num_str = x_num_str[:-3]

        x_f += 1

    if x_rstr == "":
        x_rstr = "No Rupees"
    else:
        x_rstr = "Rupees " + x_rstr

    if x_rstr_paisas != "":
        x_rstr_paisas = " and " + x_rstr_paisas + " Paisas"

    return x_rstr + x_rstr_paisas + " Only"



def generate_pdf_view(request):
    numberofrow = 0
    
    if request.method == 'POST':
        pdfdata_form = PDFForm(request.POST, request.FILES)
        if pdfdata_form.is_valid():
            pdf_data = pdfdata_form.save()
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=A4)
            width, height = A4
            styles = getSampleStyleSheet()
            styles['BodyText'].fontSize = 12
            styles['BodyText'].leading = 20
            styles['BodyText'].alignment = TA_JUSTIFY
            # words = rupee_format(pdf_data.amount)

            if pdfdata_form.cleaned_data.get('include_selfdeclarationletter'):

                # Draw letterhead image at the top, if available
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)

                # Add the address and content text
                elements = []

                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """
                elements.append(Paragraph(address, styles['BodyText']))

                # Declaration Title
                elements.append(Paragraph("<u><b>Self Declaration Letter</b></u>", styles['Title']))

                # Body Content
                body_text = """
                    Dear Sir/Madam,<br/>
                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;This refers to you that related to above subject with respect to the procurement procedure of 
                    the organization. We Perina medical suppliers  declare that we are eligible to participate 
                    in the bidding process.<br/> We have no conflict of interest with Bidder with respect to the proposed 
                    bid procurement proceedings and have no pending Litigation with the bidder.<br/><br/>
                    Thank you, <br/> Sincerely yours,
                """
                elements.append(Paragraph(body_text, styles['BodyText']))

                y_position = height - 160  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 60, width=300, height=80)
                    y_position -= 60  # Adjust y position below the signature image
                else:
                    y_position -= 60
    
                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()  # Move to the next page if needed

            if pdfdata_form.cleaned_data.get('include_warrantyletter'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Warranty/Commitment Letter</b></u>", styles['Title']))

                commitment_text = f"""
                    Dear Sir/Madam,<br/>
                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;This refers to your Tender for <b>The Procurement of {pdf_data.subject}.</b> We, Perina medical suppliers , declare that we are eligible to participate 
                    in the bidding process.if awarded for the tender we warrant that the Goods are new, unused &amp; of
                    the most recent or current models &amp; that they incorporate all recent improvements in design &amp; materials
                    unless provided otherwise in the Contract.<br/><br/>
                    We further warrant that the Goods shall be free from defects arising from any act or commission of us or
                    arising from design, materials &amp; workmanship under normal use in the conditions prevailing in the country
                    of final destination.<br/><br/>
                    The comprehensive warranty &amp; service warranty for offered product shall remain valid for the required
                    periods after the Goods or any portion thereof as the case may be, have been delivered to &amp; accepted at the
                    final destination.<br/><br/>
                    We further ensured that during the warranty period we will provide Preventive Maintenance (PPM) along
                    with corrective / breakdown maintenance whenever required.<br/><br/>
                    Thank you,<br/> Sincerely yours,<br/>
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 160  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 60, width=300, height=80)
                    y_position -= 60  # Adjust y position below the signature image
                else:
                    y_position -= 60

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()
            

            if pdfdata_form.cleaned_data.get('include_manufactureletter'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Date of Manufacture and Brand New Machine</b></u>", styles['Title']))

                commitment_text = f"""
                    Dear Sir/Madam,<br/>
                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;This refers to your tender for <b> The Procurement of {pdf_data.subject}</b>. We, Perina medical suppliers , declare that the offered equipment shall
                    be brand new manufactured after placement of the order. <br/><br/>             
                    We shall also submit certification with the date of manufacture of the machine along with shipment
                    of the system..<br/><br/>
                    Thank you, <br/>Sincerely yours,<br/>
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 160  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()


            if pdfdata_form.cleaned_data.get('include_deliverycommitmentletter'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Delivery commitment letter</b></u>", styles['Title']))

                commitment_text = f"""
                    Dear Sir/Madam,<br/>
                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;This refers to your tender for <b> The Procurement of {pdf_data.subject}</b>. We, Perina medical suppliers , declare that We Fine Surgicals Nepal Pvt.Ltd., if we are awarded the
                    contract, the Machine can be delivered within <b>{pdf_data.days} days</b> after conformation.<br/><br/>
                    Thank you, <br/> Sincerely yours,<br/>
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 160  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()


            # installation 

            if pdfdata_form.cleaned_data.get('include_installationletter'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Installation Demonstration And Training </b></u>", styles['Title']))

                commitment_text = f"""
                    Dear Sir/Madam,<br/>
                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;This refers to your tender for <b> The Procurement of {pdf_data.subject}</b>. We, Perina medical suppliers  If we are awarded the contract installation &amp;
                     demonstration of the equipment shall be carried out by the trained engineers.<br/>They shall also impart training to the engineer, technicians &amp; users at site during Installation of the system.<br/>
                Thank you, <br/> Sincerely yours<br/>
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 160  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()

            if pdfdata_form.cleaned_data.get('include_attorneyletter'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Power of attorney </b></u>", styles['Title']))

                commitment_text = f"""
                    Dear Sir/madam,<br/>
                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;We, Perina medical suppliers  Here by to assign {pdf_data.proprietor_name} {pdf_data.ourdesignation} of this company to handle all tender and quotation related to <b> The Procurement of {pdf_data.subject}</b>.<br/> He is authorized to do all necessary documentation on our behalf and purpose.<br/> 
                  Thank you, <br/> Sincerely yours
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 160  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.attornitysign:
                    attornitysign_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(attornitysign_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, f"{pdf_data.attornity_name}")
                p.drawString(50, y_position - 40, f"{pdf_data.attornity_designation}")
                p.showPage()

            if pdfdata_form.cleaned_data.get('include_bidsubmissionform'):
                
                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)

                y_position = height - 160 
                
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 120, "Date: " + date)
                    
                if pdf_data.id_no:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 140, pdf_data.id_no )

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Bid submission Forms </b></u>", styles['Title']))

                commitment_text = f"""
                    we,the undersigned declare that,<br/>
                    (a) We have examined &amp; have no reservations to the Bidding Documents, including Addenda No.: No addenda.<br/><br/>
                    (b) We offer to supply in conformity with the Bidding Documents &amp; in accordance with the Delivery Schedules
                    specified in the Schedule of Requirements the following Goods &amp; related Services:{pdf_data.subject} . The total
                    price of our Bid, excluding any discounts offered in item (d) below, is; <b> {words} (Nrs. only)</b>
                    excluding Tax and Vat.<br/><br/>
                    (c) The discounts offered and the methodology for their application are: Not applicable.<br/><br/>
                    (d) Our bid shall be valid for the period of {pdf_data.bidvaliditydays} days from the date fixed for the bid submission deadline in
                    accordance with the Bidding Document, &amp; it shall remain binding upon us &amp; may be accepted at any time
                    before the expiration of that period<br/><br/>
                    (e) If our Bid is accepted, we commit to obtain a Performance Security in the amount as specified in 1TB 41 for
                    the due performance of the Contract<br/><br/>
                    (f) We are not participating, as Bidders, in more than one Bid in this bidding process, other than alternative
                    offers in accordance with the Biding Document.<br/><br/>
                    
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))
                
                
                y_position = height - 200  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20
                
                
                p.showPage()

            if pdfdata_form.cleaned_data.get('include_bidsubmissionform'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)



                elements = []


                commitment_text = f"""
                    (g) The following commissions, gratuities, or fees, if any, have been paid or are to be paid with respect to the
                    bidding process or execution of the Contract:<br/><br/>

                   <b> Name of Receipt &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; Address &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; Reason &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;

                    Amount<br/><br/><br/>

                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; NONE<br/><br/></b>

                    (h) We understand that this bid, together with your written acceptance thereof included in your notification of
                    Ward shall constitute a binding contract between us, until a formal contract is prepared &amp; executed.<br/><br/>
                    (i) We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may
                    receive. <br/><br/>
                    (j) We declare that we are have not been blacklisted as per 1TB 3.4 and no conflict of interest in the proposed
                    procurement proceedings &amp; we have not been punished for an offence relating to the concerned profession
                    of business.<br/><br/>
                    (k) We agree to permit GoN/DP or its representative to inspect our accounts &amp; records &amp; other documents
                    relating to the bid submission &amp; to have them audited by auditors appointed by the GoN/DP.<br/><br/>
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 120  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()

# letter of technicxal bid

            if pdfdata_form.cleaned_data.get('include_technicalbid'):

                y_position = height - 200 
                
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 160, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Letter of Technical Bid </b></u>", styles['Title']))

                commitment_text = f"""
                    (a) We have examined and have no reservations to the Bidding Document, including Addenda issued in
                    accordance with Instructions to Bidders (ITB) Clause 9;<br/><br/>
                    (b) We offer to supply in conformity with the Bidding Document and in accordance with the delivery
                    schedule specified in the Section V (Schedule of Requirements), the following Goods and Related
                    Services: <b>The Procurement of {pdf_data.subject}.</b><br/><br/>
                    (c) Our Bid consisting of the Technical Bid and the Price Bid shall be valid for a period of 90 days from the
                    date fixed for the bid submission deadline in accordance with the Bidding Document, and it shall
                    remain binding upon us and may be accepted at any time before the expiration of that period;<br/><br/>
                    (d) Our firm, including any subcontractors or suppliers for any part of the Contract, has nationalities from
                    eligible countries in accordance with ITB 4.8 and meets the requirements of ITB 3.4 &amp; 3.5;<br/><br/>
                    (e) We are not participating, as a Bidder or as a subcontractor/supplier, in more than one Bid in this bidding
                    process in accordance with ITB 4.3(e), other than alternative Bids in accordance with ITB 14;<br/><br/>
                    (f) Our firm, its affiliates or subsidiaries, including any Subcontractors or Suppliers for any part of the
                    contract, has not been declared ineligible by DP, under the Purchaser’s country laws or official
                    regulations or by an act of compliance with a decision of the United Nations Security Council;<br/>
                    
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))
                
                
                y_position = height - 200  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20
                
                
                p.showPage()

            if pdfdata_form.cleaned_data.get('include_technicalbid'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)



                elements = []


                commitment_text = f"""
                    (g) We are not a government owned entity/we are a government owned entity but meet the requirements of
                    ITB 4.5;<br/><br/>
                    (h) We declare that, we including any subcontractors or suppliers for any part of the contract do not have
                    any conflict of interest in accordance with ITB 4.3 and we have not been punished for an offense
                    relating to the concerned profession or business.<br/><br/>

                    (i) The following commissions, gratuities, or fees, if any, have been paid or are to be paid with respect to the
                    bidding process or execution of the Contract:<br/><br/>

                   <b> Name of Receipt &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; Address &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; Reason &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;

                    Amount<br/><br/><br/>

                    &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; &#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; NONE<br/><br/></b>

                    (j) We declare that we are solely responsible for the authenticity of the documents submitted by us. The
                        document and information submitted by us are true and correct. If any document/information given is
                        found to be concealed at a later date, we shall accept any legal actions by the purchaser.<br/><br/>
                    (k) We agree to permit GoN/DP or its representative to inspect our accounts and records and other
                        documents relating to the bid submission and to have them audited by auditors appointed by the
                        GoN/DP.<br/><br/>
                        Thank you,<br/>Sincerely yours
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 120  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, pdf_data.proprietor_name)
                p.drawString(50, y_position - 40, pdf_data.ourdesignation)
                p.showPage()
# LETTER of price bid
            if pdfdata_form.cleaned_data.get('include_pricebid'):
                vats_amt = round(pdf_data.amount * 1.13, 2)
                vat_words = rupee_format(int(vats_amt))
                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)
                y_position = height - 200 
                
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 160, "Date: " + date)

            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Letter of Price Bid </b></u>", styles['Title']))

                commitment_text = f"""
                    (a) We have examined and have no reservations to the Bidding Document, including Addenda issued in
                    accordance with Instructions to Bidders (ITB) Clause 9;<br/><br/>
                    (b) We offer to supply in conformity with the Bidding Document and in accordance with the delivery
                    schedule specified in the Section V (Schedule of Requirements), the following Goods and Related
                    Services: <b>The Procurement of {pdf_data.subject}.</b><br/><br/>
                    (c) The total price of our Bid, excluding any discounts offered in item (d) below, is: {pdf_data.amount}(In Words {vat_words})
                    (d) The discounts offered and the methodology for their application are:
                    - The discounts offered are  NONE<br/>
                    - The exact method of calculations to determine the net price after application of discounts is shown below
                     .NONE<br/><br/>
                    (e) Our bid shall be valid for a period of {pdf_data.days} days from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period.<br/><br/>
                    (f) If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document.<br/><br/>
                    (g) We understand that this bid, together with your written acceptance thereof included in your notification of award, shall constitute a binding contract between us, until a formal contract is prepared and executed.<br/><br/>
                                        
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))
                
                
                y_position = height - 200  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20
                
                
                p.showPage()

            if pdfdata_form.cleaned_data.get('include_pricebid'):

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)



                elements = []


                commitment_text = f"""
                    (h) We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may receive. <br/><br/>
                    (i) We agree to permit the Employer/DP 1 or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors 2 appointed by the Employer. <br/><br/>
                    (j) We confirm and stand by our commitments and other declarations made in connection with the submission of our Letter of Technical Bid.<br/><br/>
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))

                y_position = height - 120  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20  # Adjust for the next element

                p.drawString(50, y_position - 20, f"Name:{pdf_data.proprietor_name}")
                p.drawString(50, y_position - 40, f"In the capacity of :{pdf_data.ourdesignation}")
                p.drawString(50, y_position - 60, "signed")
                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                    y_position -= 100  # Adjust y position below the signature image
                else:
                    y_position -=100  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, "Duly authorized to sign the bid for & on behalf of: Perina medical suppliers ")
                p.drawString(50, y_position - 40, f"Date:{pdf_data.date}")
                p.showPage()

# letter of quotation and price schedule
            if pdfdata_form.cleaned_data.get('include_quotationandprice'):
                
                vats_amt = round(pdf_data.amount * 1.13, 2)
                vat_words = rupee_format(int(vats_amt))
                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)
                
                y_position = height - 200 
                
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    p.drawImage(letterhead_path, 10, height - 100, width=600, height=100)

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    p.drawString(400, height - 110, "Date: " + date)
                   
            # Warranty commitment letter page
                # Address block as a Paragraph
                address = f"""
                    To,<br/>
                    {pdf_data.designation or ""}<br/>
                    {pdf_data.institution_name or ""}<br/>
                    {pdf_data.address or ""}
                """

                elements = []

                # Creating the paragraphs for warranty commitment
                elements.append(Paragraph(address, styles['BodyText']))

                elements.append(Paragraph("<u><b>Quotation and Price Schedule</b></u>", styles['Title']))

                commitment_text = f"""
                    Having examined the Sealed Quotation (SQ) documents, we the undersigned, offer supply and delivery of
                    {pdf_data.subject} in conformity with the said SQ documents for the sum of or
                    such other sums <u><b>{vats_amt}.</b></u> (In Words <u><b> {vat_words})</b></u>as may be ascertained in accordance with the Schedule of
                    Prices attached herewith and made part of this SQ.
                    We undertake, if our SQ is accepted, to deliver the goods in accordance with the delivery schedule specified
                    in the Schedule of Requirements.<br/><br/>
                    If our SQ is accepted, we will obtain the guarantee of bank if mentioned in contract due performance of the Contract, in the form prescribed by the Purchaser.<br/>
                    We agree to abide by this SQ for a Period of 45 days from the date fixed for SQ opening it shall remain
                    binding upon us and may be accepted at any time before the expiration of that period.<br/>
                    Until a formal Contract is prepared and executed, this SQ, together with your written acceptance thereof
                    and your notification of award, shall constitute a binding Contract between us.<br/>
                    We understand that you are not bound to accept the lowest or any SQ you may receive.<br/>
                    Dated this <u> {pdf_data.datedthis}</u><br/>Name:{pdf_data.proprietor_name}<br/>In the capacity of :{pdf_data.ourdesignation}<br/>Signed
                    
                """
                elements.append(Paragraph(commitment_text, styles['BodyText']))
                
                
                y_position = height - 120  # Start position
                for element in elements:
                    w, h = element.wrap(width - 100, height)
                    element.drawOn(p, 50, y_position - h)
                    y_position -= h + 20
                
                # Add signature at a specific position if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    p.drawImage(signature_path, 50, y_position - 60, width=300, height=80)
                    y_position -= 80  # Adjust y position below the signature image
                else:
                    y_position -=80  # Adjust y position below the signature image

                line_y_position = y_position - 5  # Adjust this value for spacing below the image
                p.setLineWidth(1)  # Set the line width if needed
                p.line(40, line_y_position, width-300, line_y_position)

                p.drawString(50, y_position - 20, "Duly authorized to sign the bid for & on behalf of: Perina medical suppliers ")
                p.drawString(50, y_position - 40, f"Date:{pdf_data.date}")
                p.showPage()
                
    


            p.save()

            # Return the PDF as a response
            buffer.seek(0)
            return HttpResponse(buffer, content_type='application/pdf', headers={'Content-Disposition': f'attachment; filename="{pdf_data.subject}.pdf" ', 'rowcount':str(numberofrow)})

    else:
        pdfdata_form = PDFForm()

    return render(request, 'pdfgenerator/forms.html', {'pdfdata_form': pdfdata_form})


def generate_word(request):
    if request.method == 'POST':
        pdfdata_form = PDFForm(request.POST, request.FILES)
        if pdfdata_form.is_valid():
            pdf_data = pdfdata_form.save()
            document = Document()
            document.sections[0].left_margin = Inches(1)
            document.sections[0].top_margin = Inches(0.2)
            document.sections[0].right_margin = Inches(0.5)
            document.sections[0].bottom_margin = Inches(0.5)
            # Helper function to validate and add image
            def add_valid_picture(doc, path, width):
                try:
                    if os.path.exists(path):
                        doc.add_picture(path, width=width)
                    else:
                        print(f"File not found: {path}")
                except Exception as e:
                    print(f"Error adding image: {e}")

            # Self Declaration Letter
            if pdfdata_form.cleaned_data.get('include_selfdeclarationletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                # document.add_paragraph("\n<u><b>Self Declaration Letter</b></u>", style='Title')
                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"self declaration letter")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    "Dear Sir/Madam,\n\n"
                    "This refers to you that related to the above subject with respect to the procurement procedure. "
                    "We Sharvil Energy Pvt. Ltd. declare that we are eligible to participate in the bidding process."
                    "We have no conflict of interest and no pending litigation.\n\nThank you,\n\nSincerely yours," 
                )
                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"_________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")

                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))

                # Add a page break after the self-declaration letter
                document.add_page_break()

            # Warranty Letter
            if pdfdata_form.cleaned_data.get('include_warrantyletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Warranty/Commitment Letter")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    f"Dear Sir/Madam,\n"
                    f"This refers to your Tender for <b>The Procurement of {pdf_data.subject}.</b> We, Perina medical suppliers , declare that we are eligible to participate "
                    f"in the bidding process. If awarded for the tender, we warrant that the Goods are new, unused & of "
                    f"the most recent or current models and that they incorporate all recent improvements in design and materials "
                    f"unless provided otherwise in the Contract.\n\n"
                    f"We further warrant that the Goods shall be free from defects arising from any act or commission of us or "
                    f"arising from design, materials and workmanship under normal use in the conditions prevailing in the country "
                    f"of final destination.\n\n"
                    f"The comprehensive warranty and service warranty for the offered product shall remain valid for the required "
                    f"periods after the Goods or any portion thereof, as the case may be, have been delivered to and accepted at the "
                    f"final destination.\n\n"
                    f"We further ensure that during the warranty period, we will provide Preventive Maintenance (PPM) along "
                    f"with corrective/breakdown maintenance whenever required.\n\n"
                    f"Thank you,\nSincerely yours,\n"
                )
                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")

                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(7.0))
                document.add_page_break()

            if pdfdata_form.cleaned_data.get('include_manufactureletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"self declaration letter")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    f"Dear Sir/Madam,\n"
                    f"This refers to your Tender for <b>The Procurement of {pdf_data.subject}.</b> We, Perina medical suppliers , declare that we are eligible to participate "
                    f"in the bidding process. If awarded for the tender, we warrant that the Goods are new, unused & of "
                    f"the most recent or current models and that they incorporate all recent improvements in design and materials "
                    f"unless provided otherwise in the Contract.\n\n"
                    f"We further warrant that the Goods shall be free from defects arising from any act or commission of us or "
                    f"arising from design, materials and workmanship under normal use in the conditions prevailing in the country "
                    f"of final destination.\n\n"
                    f"The comprehensive warranty and service warranty for the offered product shall remain valid for the required "
                    f"periods after the Goods or any portion thereof, as the case may be, have been delivered to and accepted at the "
                    f"final destination.\n\n"
                    f"We further ensure that during the warranty period, we will provide Preventive Maintenance (PPM) along "
                    f"with corrective/breakdown maintenance whenever required.\n\n"
                    f"Thank you,\nSincerely yours,\n"
                )
                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")

                document.add_page_break()


            if pdfdata_form.cleaned_data.get('include_manufactureletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Date of Manufacture and Brand New Machine")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    "Dear Sir/Madam,\n\n"
                    f"        This refers to your tender for **The Procurement of {pdf_data.subject}**. We, Sharvil Energy Pvt. Ltd., "
                    "declare that the offered equipment shall be brand new, manufactured after the placement of the order.\n\n"
                    "We shall also submit certification with the date of manufacture of the machine along with the shipment "
                    "of the system.\n\n"
                    "Thank you,\nSincerely yours,\n"
                )

                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")

                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))

                document.add_page_break()



            if pdfdata_form.cleaned_data.get('include_deliverycommitmentletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Delivery commitment letter")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    "Dear Sir/Madam,\n\n"
                    f"        This refers to your tender for **The Procurement of {pdf_data.subject}**. "
                    "We, Sharvil Energy Pvt. Ltd., declare that if we are awarded the "
                    f"contract, the Machine can be delivered within **{pdf_data.days} days** after confirmation.\n\n"
                    "Thank you,\nSincerely yours,\n"
                )

                document.add_paragraph(body_text)
                
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")
                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))
                document.add_page_break()


            if pdfdata_form.cleaned_data.get('include_installationletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Installation Demonstration And Training")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    f"Dear Sir/Madam,\n\n"
                    f"        This refers to your tender for **The Procurement of {pdf_data.subject}**. "
                    "We, Sharvil Energy Pvt. Ltd., declare that if we are awarded the contract, "
                    "the installation and demonstration of the equipment shall be carried out by trained engineers.\n"
                    "They shall also impart training to the engineers, technicians, and users at the site during the installation of the system.\n\n"
                    "Thank you,\nSincerely yours\n"
)
                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")
                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))

                document.add_page_break()



            if pdfdata_form.cleaned_data.get('include_attorneyletter'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Power of attorney")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    f"Dear Sir/Madam,\n\n"
                    f"        We, Sharvil Energy Pvt. Ltd., hereby assign {pdf_data.proprietor_name}, {pdf_data.ourdesignation} of this company, "
                    f"to handle all tenders and quotations related to **The Procurement of {pdf_data.subject}**.\n"
                    "He is authorized to do all necessary documentation on our behalf and purpose.\n\n"
                    "Thank you,\nSincerely yours\n"
                )

                document.add_paragraph(body_text)

                if pdf_data.attornitysign:
                    attornitysign_path = os.path.join(settings.MEDIA_ROOT, pdf_data.attornitysign.name)
                    add_valid_picture(document, attornitysign_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.attornity_name}\n{pdf_data.attornity_designation}")
                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))
                document.add_page_break()



            if pdfdata_form.cleaned_data.get('include_bidsubmissionform'):

                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Bid submission Forms")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    f"We, the undersigned, declare that:\n\n"
                    "(a) We have examined and have no reservations to the Bidding Documents, including Addenda No.: No addenda.\n\n"
                    "(b) We offer to supply in conformity with the Bidding Documents and in accordance with the Delivery Schedules "
                    f"specified in the Schedule of Requirements the following Goods and related Services: {pdf_data.subject}. "
                    f"The total price of our Bid, excluding any discounts offered in item (c) below, is: {words} (Nrs) excluding Tax and VAT.\n\n"
                    "(c) The discounts offered and the methodology for their application are: Not applicable.\n\n"
                    f"(d) Our bid shall be valid for the period of {pdf_data.bidvaliditydays} days from the date fixed for the bid submission deadline in "
                    "accordance with the Bidding Document, and it shall remain binding upon us and may be accepted at any time "
                    "before the expiration of that period.\n\n"
                    "(e) If our Bid is accepted, we commit to obtain a Performance Security in the amount specified in ITB 41 for "
                    "the due performance of the Contract.\n\n"
                    "(f) We are not participating, as Bidders, in more than one Bid in this bidding process, other than alternative "
                    "offers in accordance with the Bidding Document.\n\n"
                )

                document.add_paragraph(body_text)

                # if pdf_data.signature:
                #     signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                #     add_valid_picture(document, signature_path, width=Inches(2.0))
                #     document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")
                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))
                document.add_page_break()

            if pdfdata_form.cleaned_data.get('include_bidsubmissionform'):

                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))


                body_text = (
                "(g) The following commissions, gratuities, or fees, if any, have been paid or are to be paid with respect to the "
                "bidding process or execution of the Contract:\n\n"
                "Name of Recipient                  Address                    Reason                  Amount\n"
                "------------------------------------------------------------------------------------------\n"
                "                                                NONE\n\n"
                "(h) We understand that this bid, together with your written acceptance thereof included in your notification of "
                "award shall constitute a binding contract between us, until a formal contract is prepared and executed.\n\n"
                "(i) We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may "
                "receive.\n\n"
                "(j) We declare that we have not been blacklisted as per ITB 3.4 and that there is no conflict of interest in the proposed "
                "procurement proceedings. We further declare that we have not been punished for any offence relating to the "
                "concerned profession or business.\n\n"
                "(k) We agree to permit GoN/DP or its representative to inspect our accounts, records, and other documents "
                "relating to the bid submission and to have them audited by auditors appointed by the GoN/DP.\n\n"
            )

                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")
                document.add_page_break()

            if pdfdata_form.cleaned_data.get('include_technicalbid'):

                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Letter of Technical Bid")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    "(a) We have examined and have no reservations to the Bidding Document, including Addenda issued in "
                    "accordance with Instructions to Bidders (ITB) Clause 9.\n\n"
                    f"(b) We offer to supply in conformity with the Bidding Document and in accordance with the delivery "
                    "schedule specified in Section V (Schedule of Requirements), the following Goods and Related "
                    f"Services: The Procurement of {pdf_data.subject}.\n\n"
                    "(c) Our Bid consisting of the Technical Bid and the Price Bid shall be valid for a period of 90 days from the "
                    "date fixed for the bid submission deadline in accordance with the Bidding Document, and it shall "
                    "remain binding upon us and may be accepted at any time before the expiration of that period.\n\n"
                    "(d) Our firm, including any subcontractors or suppliers for any part of the Contract, has nationalities from "
                    "eligible countries in accordance with ITB 4.8 and meets the requirements of ITB 3.4 and 3.5.\n\n"
                    "(e) We are not participating, as a Bidder or as a subcontractor/supplier, in more than one Bid in this bidding "
                    "process in accordance with ITB 4.3(e), other than alternative Bids in accordance with ITB 14.\n\n"
                    "(f) Our firm, its affiliates or subsidiaries, including any Subcontractors or Suppliers for any part of the "
                    "contract, has not been declared ineligible by DP, under the Purchaser’s country laws or official "
                    "regulations or by an act of compliance with a decision of the United Nations Security Council.\n\n"
                )

                document.add_paragraph(body_text)

                # if pdf_data.signature:
                #     signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                #     add_valid_picture(document, signature_path, width=Inches(2.0))
                #     document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")
                # if pdf_data.footer:
                #     footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
                #     add_valid_picture(document, footer_path, width=Inches(6.0))
                document.add_page_break()

            if pdfdata_form.cleaned_data.get('include_technicalbid'):

                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))


                body_text = (
                    "(g) We are not a government-owned entity / we are a government-owned entity but meet the requirements of "
                    "ITB 4.5.\n\n"
                    "(h) We declare that we, including any subcontractors or suppliers for any part of the contract, do not have "
                    "any conflict of interest in accordance with ITB 4.3 and have not been punished for an offense "
                    "relating to the concerned profession or business.\n\n"
                    "(i) The following commissions, gratuities, or fees, if any, have been paid or are to be paid with respect to the "
                    "bidding process or execution of the contract:\n\n"
                    "Name of Recipient                 Address                        Reason                              Amount\n"
                    "------------------------------------------------------------------------------------------------------------\n"
                    "                                                    NONE\n\n"
                    "(j) We declare that we are solely responsible for the authenticity of the documents submitted by us. The "
                    "documents and information submitted are true and correct. If any document or information is "
                    "found to be concealed or misrepresented at a later date, we shall accept any legal action by the purchaser.\n\n"
                    "(k) We agree to permit GoN/DP or its representative to inspect our accounts, records, and other "
                    "documents relating to the bid submission, and to have them audited by auditors appointed by the "
                    "GoN/DP.\n\n"
                    "Thank you,\nSincerely yours"
                )

                document.add_paragraph(body_text)

                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    add_valid_picture(document, signature_path, width=Inches(2.0))
                    document.add_paragraph(f"________________________\n{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")

                document.add_page_break()

            if pdfdata_form.cleaned_data.get('include_pricebid'):
                vats_amt = "{:,.2f}".format(pdf_data.amount * 1.13)
                vat_words = rupee_format(int(float(vats_amt.replace(',', ''))))
                print("Formatted Words:change", vat_words)
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"Letter of Price Bid")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    "(a) We have examined and have no reservations to the Bidding Document, including Addenda issued in "
                    "accordance with Instructions to Bidders (ITB) Clause 9.\n\n"
                    f"(b) We offer to supply in conformity with the Bidding Document and in accordance with the delivery "
                    f"schedule specified in Section V (Schedule of Requirements), the following Goods and Related "
                    f"Services: The Procurement of {pdf_data.subject}.\n\n"
                    f"(c) The total price of our Bid, excluding any discounts offered in item (d) below, is: {vats_amt} "
                    f"(In Words {vat_words}).\n\n"
                    "(d) The discounts offered and the methodology for their application are:\n"
                    "- The discounts offered are NONE\n"
                    "- The exact method of calculations to determine the net price after application of discounts is shown below: NONE\n\n"
                    f"(e) Our bid shall be valid for a period of {pdf_data.bidvaliditydays} days from the date fixed for the bid submission deadline "
                    "in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time "
                    "before the expiration of that period.\n\n"
                    "(f) If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document.\n\n"
                )

                document.add_paragraph(body_text)

                
                document.add_page_break()


            if pdfdata_form.cleaned_data.get('include_pricebid'):

                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)

                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))


                body_text = (
                    f"(g) We understand that this bid, together with your written acceptance thereof included in your notification of award, shall constitute a binding contract between us, until a formal contract is prepared and executed.\n\n"
                    f"(h) We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may receive. \n\n"
                    f"(i) We agree to permit the Employer/DP 1 or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors 2 appointed by the Employer. \n\n"
                    f"(j) We confirm and stand by our commitments and other declarations made in connection with the submission of our Letter of Technical Bid.\n\n"
                )
                document.add_paragraph(body_text)

                # Add bidder information
                document.add_paragraph(f"Name: {pdf_data.proprietor_name}")
                document.add_paragraph(f"In the capacity of: {pdf_data.ourdesignation}")
                document.add_paragraph("Signed")

                # Add signature if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    document.add_picture(signature_path, width=Pt(150))

                # Add company information
                document.add_paragraph("Duly authorized to sign the bid for & on behalf of: Perina medical suppliers ")
                document.add_paragraph(f"Date: {pdf_data.date}")


               

                document.add_page_break()


            if pdfdata_form.cleaned_data.get('include_quotationandprice'):
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    add_valid_picture(document, letterhead_path, width=Inches(7.0))

                if pdf_data.date:
                    date = pdf_data.date.strftime("%Y-%m-%d")
                    # document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}")
                    subject_paragraph = document.add_paragraph()
                    run = subject_paragraph.add_run(f"Date: {date}\n{pdf_data.id_no}")
                    run.font.size = Pt(12)  # Adjust font size as needed
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                if pdf_data.designation:
                    document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
                else:
                    document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

                subject_paragraph = document.add_paragraph()
                run = subject_paragraph.add_run(f"quotation and price schedule")
                run.font.size = Pt(14)  # Adjust font size as needed
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                body_text = (
                    
                    "Having examined the Sealed Quotation (SQ) documents, we the undersigned, offer supply and delivery of "
                    f"{pdf_data.subject} in conformity with the said SQ documents for the sum of or "
                    f"such other sums {vats_amt}. (In Words {vat_words}) as may be ascertained in accordance with the Schedule of "
                    "Prices attached herewith and made part of this SQ.\n"
                    "We undertake, if our SQ is accepted, to deliver the goods in accordance with the delivery schedule specified "
                    "in the Schedule of Requirements.\n"
                    "If our SQ is accepted, we will obtain the guarantee of bank if mentioned in contract due performance of the Contract, in the form prescribed by the Purchaser.\n"
                    f"We agree to abide by this SQ for a Period of {pdf_data.bidvaliditydays} days from the date fixed for SQ opening it shall remain "
                    "binding upon us and may be accepted at any time before the expiration of that period.\n"
                    "Until a formal Contract is prepared and executed, this SQ, together with your written acceptance thereof "
                    "and your notification of award, shall constitute a binding Contract between us.\n"
                    "We understand that you are not bound to accept the lowest or any SQ you may receive.\n"
                    f"Dated this  {pdf_data.datedthis}"
                    
                )

                document.add_paragraph(body_text)

                 # Add bidder information
                document.add_paragraph(f"Name: {pdf_data.proprietor_name}")
                document.add_paragraph(f"In the capacity of: {pdf_data.ourdesignation}")
                document.add_paragraph("Signed")

                # Add signature if available
                if pdf_data.signature:
                    signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                    document.add_picture(signature_path, width=Pt(150))

                # Add company information
                document.add_paragraph("Duly authorized to sign the bid for & on behalf of: Perina medical suppliers ")
                document.add_paragraph(f"Date: {pdf_data.date}")


                
                document.add_page_break()
            

            # Prepare response
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)

            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=self_declaration.docx'
            return response



def calculate_column_widths(data, font_name="Helvetica", font_size=10, min_width=50, max_width=200):
    column_widths = []
    for col in data.columns:
        max_cell_width = max(
            stringWidth(str(cell), font_name, font_size) if cell else 0
            for cell in data[col]
        )
        column_widths.append(min(max(max_cell_width + 10, min_width), max_width))
    return column_widths





def generate_priceschedule(request):
    if request.method == 'POST':
        pricescheduleform_form = pricescheduleform(request.POST, request.FILES)
        if pricescheduleform_form.is_valid():
            pdf_data = pricescheduleform_form.save()
            buffer = BytesIO()
            styles = getSampleStyleSheet()
            centered_heading_style = styles['Heading3']
            centered_heading_style.alignment = TA_CENTER
            styles['BodyText'].fontSize = 10
            styles['BodyText'].leading = 20
            width, height = A4

            # Prepare the content for the document
            story = []

            # Function to add header/footer
            def add_header_footer(canvas, doc):
                canvas.saveState()
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    canvas.drawImage(letterhead_path, 40, height - 100, width=500, height=80)
                canvas.restoreState()

            # Add "Price Schedule for Goods" heading
            heading = Paragraph(f"<u>Price Schedule for Goods</u> <br/><br/> Name of bidder : {pdf_data.institution_name}", centered_heading_style)
            story.append(heading)
            story.append(Spacer(1, 20))

            # Handle Excel table data
            if pdf_data.excel:
                excel_path = os.path.join(settings.MEDIA_ROOT, pdf_data.excel.name)
                data = pd.read_excel(excel_path, engine='openpyxl')
                data = data.where(pd.notnull(data), None)  # Handle NaN values
                data_list = [data.columns.tolist()] + data.values.tolist()
                vats_amt = (pdf_data.amount*1.13)
                vat_words = rupee_format(int(vats_amt))
                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)
                # Prepare the table data with Paragraphs
                table_data = [
                    [Paragraph(str(cell) if cell else '', styles["BodyText"]) for cell in row]
                    for row in data_list
                ]

                # Calculate column widths
                # Calculate column widths dynamically
                column_widths = calculate_column_widths(data, font_name="Helvetica", font_size=10)

                # Normalize widths to fit within the page
                available_width = width - 60  # Account for page margins
                total_width = sum(column_widths)
                if total_width > available_width:
                    scaling_factor = available_width / total_width
                    column_widths = [w * scaling_factor for w in column_widths]

                # Prepare table data
                table_data = [
                    [Paragraph(str(cell) if cell else '', styles["BodyText"]) for cell in row]
                    for row in data_list
                ]

                # Create table
                table = Table(table_data, colWidths=column_widths)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)


            # Add a spacer after the table
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"In Words :{vat_words}"))
            story.append(Paragraph(f"Notes*:{pdf_data.notes}"))

            # Add signature section
            # story.append(Spacer(1, 50))
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                story.append(Image(signature_path, width=150, height=50))
                story.append(Paragraph("Authorized Signature", styles["BodyText"]))
            else:
                story.append(Paragraph("Authorized Signature: ______________________", styles["BodyText"]))

            # Build the PDF
            doc = SimpleDocTemplate(
                buffer, pagesize=A4,
                rightMargin=30, leftMargin=30,
                topMargin=100, bottomMargin=30
            )
            doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)

            # Return the PDF as a response
            buffer.seek(0)
            return HttpResponse(buffer, content_type='application/pdf', headers={'Content-Disposition': f'attachment; filename="finesurgical.pdf"'})

    else:
        pricescheduleform_form = pricescheduleform()

    return render(request, 'pdfgenerator/priceschedule.html', {'pricescheduleform_form': pricescheduleform_form})



def upload(request):
    if request.method == 'POST':
        quotation_form = quotationform(request.POST, request.FILES)
        if quotation_form.is_valid():
            pdf_data = quotation_form.save()
            buffer = BytesIO()
            styles = getSampleStyleSheet()
            width, height = A4
            story = []

            # Add PDF content (address, table, etc.)
            story.append(Paragraph("Quotation Generated", styles['Title']))

            # Load and Process Excel Data
            if pdf_data.excel:
                excel_path = os.path.join(settings.MEDIA_ROOT, pdf_data.excel.name)
                data = pd.read_excel(excel_path, engine='openpyxl')
                rows = data.iloc[:, 0].tolist()  # Assuming first column for rows

            # Generate PDF
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            doc.build(story)

            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="quotation.pdf"'

            # Pass context to template (for upload section)
            return redirect('upload_pdf', quotation_id=pdf_data.id)
    else:
        quotation_form = quotationform()

    return render(request, 'pdfgenerator/quotation.html', {'quotation_form': quotation_form})

# def generate_word(request):
#     if request.method == 'POST':
#         pdfdata_form = PDFForm(request.POST, request.FILES)
#         if pdfdata_form.is_valid():
#             pdf_data = pdfdata_form.save()
#             document = Document()

#             # Helper function to validate and add image
#             def add_valid_picture(doc, path, width):
#                 try:
#                     with Image.open(path) as img:
#                         img.verify()  # Validate image
#                         doc.add_picture(path, width=width)
#                 except (IOError, OSError) as e:
#                     print(f"Invalid image or file not found: {path}. Error: {e}")

#             # Add letterhead image
#             if pdf_data.letterhead:
#                 letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
#                 add_valid_picture(document, letterhead_path, width=Inches(6.0))

#             # Add text details
#             if pdf_data.date:
#                 date = pdf_data.date.strftime("%Y-%m-%d")
#                 document.add_paragraph(f"Date: {date}\n{pdf_data.id_no}\n{pdf_data.nameofcontract}")

#             if pdf_data.designation:
#                 document.add_paragraph(f"To,\n{pdf_data.designation}\n{pdf_data.institution_name}\n{pdf_data.address}")
#             else:
#                 document.add_paragraph(f"To,\n{pdf_data.institution_name}\n{pdf_data.address}")

#             document.add_paragraph("\n<u><b>Self Declaration Letter</b></u>", style='Title')

#             body_text = (
#                 "Dear Sir/Madam,\n"
#                 "This refers to you that related to the above subject with respect to the procurement procedure. "
#                 "We Sharvil Energy Pvt. Ltd. declare that we are eligible to participate in the bidding process. "
#                 "We have no conflict of interest and no pending litigation.\n\nThank you,\nSincerely yours," 
#             )
#             document.add_paragraph(body_text)

#             # Add signature image
#             if pdf_data.signature:
#                 signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
#                 add_valid_picture(document, signature_path, width=Inches(2.0))
#                 document.add_paragraph(f"{pdf_data.proprietor_name}\n{pdf_data.ourdesignation}")

#             # Add footer image
#             if pdf_data.footer:
#                 footer_path = os.path.join(settings.MEDIA_ROOT, pdf_data.footer.name)
#                 add_valid_picture(document, footer_path, width=Inches(6.0))

#             # Prepare response
#             buffer = BytesIO()
#             document.save(buffer)
#             buffer.seek(0)
#             response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#             response['Content-Disposition'] = 'attachment; filename=self_declaration.docx'
#             return response




def generate_quotation(request):
    if request.method == 'POST':
        quotation_form = quotationform(request.POST, request.FILES)
        if quotation_form.is_valid():
            pdf_data = quotation_form.save()
            # p = canvas.Canvas(buffer, pagesize=A4)
            buffer = BytesIO()
            styles = getSampleStyleSheet()
            centered_heading_style = styles['Heading3']
            centered_heading_style.alignment = TA_CENTER
            styles['BodyText'].fontSize = 10
            styles['BodyText'].leading = 20
            width, height = A4

            # Prepare the content for the document
            story = []

            # Function to add header/footer
            def add_header_footer(canvas, doc):
                canvas.saveState()
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    canvas.drawImage(letterhead_path, 40, height - 100, width=500, height=80)
                canvas.restoreState()

            address = f"""
                To,<br/>
                {pdf_data.designation or ""}<br/>
                {pdf_data.institution_name or ""}<br/>
                {pdf_data.address or ""}<br/>
                
        """
            story.append(Paragraph(address, styles['BodyText']))

            story.append(Paragraph(f"<u><b>subject </b></u>:{pdf_data.subject}", styles['Heading3']))
            story.append(Paragraph(f"Dear sir/madam<br/>&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; {pdf_data.paragraph}", styles['BodyText']))


            # Handle Excel table data
            if pdf_data.excel:
                excel_path = os.path.join(settings.MEDIA_ROOT, pdf_data.excel.name)
                data = pd.read_excel(excel_path, engine='openpyxl')
                data = data.where(pd.notnull(data), None)  # Handle NaN values
                data_list = [data.columns.tolist()] + data.values.tolist()
                vats_amt = (pdf_data.amount*1.13)
                vat_words = rupee_format(int(vats_amt))
                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)
                # Prepare the table data with Paragraphs
                table_data = [
                    [Paragraph(str(cell) if cell else '', styles["BodyText"]) for cell in row]
                    for row in data_list
                ]

                # Calculate column widths
                # Calculate column widths dynamically
                column_widths = calculate_column_widths(data, font_name="Helvetica", font_size=10)

                # Normalize widths to fit within the page
                available_width = width - 60  # Account for page margins
                total_width = sum(column_widths)
                if total_width > available_width:
                    scaling_factor = available_width / total_width
                    column_widths = [w * scaling_factor for w in column_widths]

                # Prepare table data
                table_data = [
                    [Paragraph(str(cell) if cell else '', styles["BodyText"]) for cell in row]
                    for row in data_list
                ]

                # Create table
                table = Table(table_data, colWidths=column_widths)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)

            # story.append(Paragraph(f"Name:{pdf_data.propriter_name}"))
            story.append(Paragraph(f"In the capacity of :{pdf_data.ourdesignation}"))
            story.drawString("signed")
            # Add signature at a specific position if available
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                story.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                y_position -= 100  # Adjust y position below the signature image
            else:
                y_position -=100  # Adjust y position below the signature image

            line_y_position = y_position - 5  # Adjust this value for spacing below the image
            story.setLineWidth(1)  # Set the line width if needed
            story.line(40, line_y_position, width-300, line_y_position)

            story.drawString( "Duly authorized to sign the bid for & on behalf of: Perina medical suppliers ")
            story.drawString(f"Date:{pdf_data.date}")

            # Add a spacer after the table
            # story.append(Spacer(1, 10))
            # story.append(Paragraph(f"In Words :{vat_words}"))
            # story.append(Paragraph(f"Notes*:{pdf_data.notes}"))

            # # Add signature section
            # story.append(Spacer(1, 20))
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                story.append(Image(signature_path, width=100, height=50, hAlign='LEFT'))

                story.append(Paragraph("Authorized Signature", styles["BodyText"]))
            else:
                story.append(Paragraph("Authorized Signature: ______________________", styles["BodyText"]))

            doc = SimpleDocTemplate(
                buffer, pagesize=A4,
                rightMargin=30, leftMargin=30,
                topMargin=100, bottomMargin=30
            )
            doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)

            # Return the PDF as a response
            buffer.seek(0)
            return HttpResponse(buffer, content_type='application/pdf', headers={'Content-Disposition': f'attachment; filename="finesurgical.pdf"'})

    else:
        quotation_form = quotationform()

    return render(request, 'pdfgenerator/quotation.html', {'quotation_form': quotation_form})
    


# word generation of quotation

def quotation_docs(request):
    if request.method == 'POST':
        quotation_form = quotationform(request.POST, request.FILES)
        if quotation_form.is_valid():
            pdf_data = quotation_form.save()
            document = Document()
            # Helper function to validate and add an image
            def add_valid_picture(doc, path, width):
                try:
                    if os.path.exists(path):
                        doc.add_picture(path, width=width)
                    else:
                        print(f"File not found: {path}")
                except Exception as e:
                    print(f"Error adding image: {e}")

            # Add letterhead image
            document.sections[0].left_margin = Inches(0.5)
            document.sections[0].top_margin = Inches(0.2)
            if pdf_data.letterhead:
                letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                add_valid_picture(document, letterhead_path, width=Inches(8))

            # Add recipient address
            address = f"""To,\n {pdf_data.institution_name or ""}\n{pdf_data.address or ""}"""
            document.add_paragraph(address)

            # Add subject
            subject_paragraph = document.add_paragraph()
            run = subject_paragraph.add_run(f"Subject: {pdf_data.subject}")
            run.font.size = Pt(14)  # Adjust font size as needed
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
            subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add introductory text
            document.add_paragraph(f"Dear Sir/Madam,\n{pdf_data.paragraph}")

            # Add table content from Excel
            if pdf_data.excel:
                excel_path = os.path.join(settings.MEDIA_ROOT, pdf_data.excel.name)
                data = pd.read_excel(excel_path, engine='openpyxl').fillna("")
                table = document.add_table(rows=1, cols=len(data.columns))
                table.style = 'Table Grid'

                # Add headers
                for i, header in enumerate(data.columns):
                    table.rows[0].cells[i].text = str(header)

                # Add rows
                for row in data.itertuples(index=False):
                    table_row = table.add_row().cells
                    for i, value in enumerate(row):
                        table_row[i].text = str(value)

            # Add totals and notes
            document.add_paragraph(f"Total Amount (incl. VAT): {pdf_data.amount * 1.13}")
            document.add_paragraph(f"Notes: {pdf_data.notes}")

            
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                add_valid_picture(document, signature_path, width=Inches(2.0))
            # document.add_paragraph(f"______________________\n{pdf_data.propritername}\n{pdf_data.designation}")

            # Prepare response
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return HttpResponse(
                buffer,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': 'attachment; filename="quotation.docx"'}
            )
    else:
        quotation_form = quotationform()

    return render(request, 'generatepdf/quotation.html', {'quotation_form': quotation_form})





def generate_quotation2(request):
    if request.method == 'POST':
        quotation_form = quotationform(request.POST, request.FILES)
        if quotation_form.is_valid():
            pdf_data = quotation_form.save()
            # p = canvas.Canvas(buffer, pagesize=A4)
            buffer = BytesIO()
            styles = getSampleStyleSheet()
            centered_heading_style = styles['Heading3']
            centered_heading_style.alignment = TA_CENTER
            styles['BodyText'].fontSize = 10
            styles['BodyText'].leading = 20
            width, height = A4

            # Prepare the content for the document
            story = []

            # Function to add header/footer
            def add_header_footer(canvas, doc):
                canvas.saveState()
                if pdf_data.letterhead:
                    letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                    canvas.drawImage(letterhead_path, 40, height - 100, width=500, height=80)
                canvas.restoreState()

            address = f"""
                To,<br/>
                {pdf_data.designation or ""}<br/>
                {pdf_data.institution_name or ""}<br/>
                {pdf_data.address or ""}<br/>
                
        """
            story.append(Paragraph(address, styles['BodyText']))

            story.append(Paragraph(f"<u><b>subject </b></u>:{pdf_data.subject}", styles['Heading3']))
            story.append(Paragraph(f"Dear sir/madam<br/>&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160;&#160; {pdf_data.paragraph}", styles['BodyText']))


            # Handle Excel table data
            if pdf_data.excel:
                excel_path = os.path.join(settings.MEDIA_ROOT, pdf_data.excel.name)
                data = pd.read_excel(excel_path, engine='openpyxl')
                data = data.where(pd.notnull(data), None)  # Handle NaN values
                data_list = [data.columns.tolist()] + data.values.tolist()
                vats_amt = (pdf_data.amount*1.13)
                vat_words = rupee_format(int(vats_amt))
                words = rupee_format(int(pdf_data.amount))
                print("Formatted Words:", words)
                # Prepare the table data with Paragraphs
                table_data = [
                    [Paragraph(str(cell) if cell else '', styles["BodyText"]) for cell in row]
                    for row in data_list
                ]

                # Calculate column widths
                # Calculate column widths dynamically
                column_widths = calculate_column_widths(data, font_name="Helvetica", font_size=10)

                # Normalize widths to fit within the page
                available_width = width - 60  # Account for page margins
                total_width = sum(column_widths)
                if total_width > available_width:
                    scaling_factor = available_width / total_width
                    column_widths = [w * scaling_factor for w in column_widths]

                # Prepare table data
                table_data = [
                    [Paragraph(str(cell) if cell else '', styles["BodyText"]) for cell in row]
                    for row in data_list
                ]

                # Create table
                table = Table(table_data, colWidths=column_widths)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 5),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)

            story.append(Paragraph(f"Name:{pdf_data.propriter_name}"))
            story.append(Paragraph(50, y_position - 40, f"In the capacity of :{pdf_data.ourdesignation}"))
            story.drawString(50, y_position - 60, "signed")
            # Add signature at a specific position if available
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                story.drawImage(signature_path, 50, y_position - 80, width=300, height=80)
                y_position -= 100  # Adjust y position below the signature image
            else:
                y_position -=100  # Adjust y position below the signature image

            line_y_position = y_position - 5  # Adjust this value for spacing below the image
            story.setLineWidth(1)  # Set the line width if needed
            story.line(40, line_y_position, width-300, line_y_position)

            story.drawString(50, y_position - 20, "Duly authorized to sign the bid for & on behalf of: Perina medical suppliers ")
            story.drawString(50, y_position - 40, f"Date:{pdf_data.date}")

            
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                story.append(Image(signature_path, width=100, height=50, hAlign='LEFT'))

                story.append(Paragraph("Authorized Signature", styles["BodyText"]))
            else:
                story.append(Paragraph("Authorized Signature: ______________________", styles["BodyText"]))

            doc = SimpleDocTemplate(
                buffer, pagesize=A4,
                rightMargin=30, leftMargin=30,
                topMargin=100, bottomMargin=30
            )
            doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)

            # Return the PDF as a response
            buffer.seek(0)
            return HttpResponse(buffer, content_type='application/pdf', headers={'Content-Disposition': f'attachment; filename="finesurgical.pdf"'})

    else:
        quotation_form = quotationform()

    return render(request, 'pdfgenerator/quotation2.html', {'quotation_form': quotation_form})
    


# word generation of quotation

def quotation_docs2(request):
    if request.method == 'POST':
        quotation_form = quotationform(request.POST, request.FILES)
        if quotation_form.is_valid():
            pdf_data = quotation_form.save()
            document = Document()
            # Helper function to validate and add an image
            def add_valid_picture(doc, path, width):
                try:
                    if os.path.exists(path):
                        doc.add_picture(path, width=width)
                    else:
                        print(f"File not found: {path}")
                except Exception as e:
                    print(f"Error adding image: {e}")

            # Add letterhead image
            document.sections[0].left_margin = Inches(0.5)
            document.sections[0].top_margin = Inches(0.2)
            if pdf_data.letterhead:
                letterhead_path = os.path.join(settings.MEDIA_ROOT, pdf_data.letterhead.name)
                add_valid_picture(document, letterhead_path, width=Inches(8))

            # Add recipient address
            address = f"""To,\n {pdf_data.institution_name or ""}\n{pdf_data.address or ""}"""
            document.add_paragraph(address)

            # Add subject
            subject_paragraph = document.add_paragraph()
            run = subject_paragraph.add_run(f"Subject: {pdf_data.subject}")
            run.font.size = Pt(14)  # Adjust font size as needed
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
            subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add introductory text
            document.add_paragraph(f"Dear Sir/Madam,\n{pdf_data.paragraph}")

            # Add table content from Excel
            if pdf_data.excel:
                excel_path = os.path.join(settings.MEDIA_ROOT, pdf_data.excel.name)
                data = pd.read_excel(excel_path, engine='openpyxl').fillna("")
                table = document.add_table(rows=1, cols=len(data.columns))
                table.style = 'Table Grid'

                # Add headers
                for i, header in enumerate(data.columns):
                    table.rows[0].cells[i].text = str(header)

                # Add rows
                for row in data.itertuples(index=False):
                    table_row = table.add_row().cells
                    for i, value in enumerate(row):
                        table_row[i].text = str(value)

            # Add totals and notes
            document.add_paragraph(f"Total Amount (incl. VAT): {pdf_data.amount * 1.13}")
            document.add_paragraph(f"Notes: {pdf_data.notes}")

            
            if pdf_data.signature:
                signature_path = os.path.join(settings.MEDIA_ROOT, pdf_data.signature.name)
                add_valid_picture(document, signature_path, width=Inches(2.0))
            # document.add_paragraph(f"______________________\n{pdf_data.propritername}\n{pdf_data.designation}")

            # Prepare response
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return HttpResponse(
                buffer,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': 'attachment; filename="quotation.docx"'}
            )
    else:
        quotation_form = quotationform()

    return render(request, 'generatepdf/quotation2.html', {'quotation_form': quotation_form})




def upload_pdfs(request, quotation_id):
    quotation = get_object_or_404(generatequotation, id=quotation_id)
    
    # Extract Excel data rows for uploading PDFs
    excel_path = os.path.join(settings.MEDIA_ROOT, quotation.excel.name)
    data = pd.read_excel(excel_path, engine='openpyxl')
    rows = data.iloc[:, 1].tolist()  # Assuming first column holds row titles
    rows = rows[:-3]  # Get the last 3 rows

    if request.method == 'POST':
        # Upload Financial Docs1 and BOQ
        financial_docs = request.FILES.get('financial_docs')
        boq_docs = request.FILES.get('boq_docs')
        vat_docs = request.FILES.get('vat_docs')  # VAT is uploaded once

        # Track uploaded files
        uploaded_files = []
        
        # Save uploaded files
        if financial_docs:
            uploaded_files.append(UploadedFile.objects.create(
                quotation=quotation,
                file=financial_docs,
                file_type='Financial Docs'
            ))
        if boq_docs:
            uploaded_files.append(UploadedFile.objects.create(
                quotation=quotation,
                file=boq_docs,
                file_type='BOQ'
            ))
        if vat_docs:
            vat_instance = UploadedFile.objects.create(
                quotation=quotation,
                file=vat_docs,
                file_type='VAT'
            )
            uploaded_files.append(vat_instance)
        
        # Upload files for each row
        for row in rows:
            print(f"Row: {row}")
            catalogue_file = request.FILES.get(f'catalogue_{row}')
            print(f"Catalogue file: {catalogue_file}")
            ce_file = request.FILES.get(f'ce_{row}')
            print(f"CE file: {ce_file}")
            iso_file = request.FILES.get(f'iso_{row}')
            print(f"ISO file: {iso_file}")
            if catalogue_file:
                try:
                    uploaded_files.append(UploadedFile.objects.create(
                        quotation=quotation,
                        file=catalogue_file,
                        row_name=row,
                        file_type='Catalogue'
                    ))
                except Exception as e:
                    print(f"Error uploading file {catalogue_file.name}: {e}")

            if ce_file:
                try:
                    uploaded_files.append(UploadedFile.objects.create(
                        quotation=quotation,
                        file=ce_file,
                        row_name=row,
                        file_type='CE'
                    ))
                except Exception as e:
                    print(f"Error uploading file {ce_file.name}: {e}")
            

            if iso_file:
                try:
                    uploaded_files.append(UploadedFile.objects.create(
                        quotation=quotation,
                        file=iso_file,
                        row_name=row,
                        file_type='ISO'
                    ))
                except Exception as e:
                    print(f"Error uploading file {iso_file.name}: {e}")
            

        # Merge all PDFs
        merger = PdfMerger()
        
        for uploaded_file in uploaded_files:
            file_path = uploaded_file.file.path
            print(f"Merging file: {file_path}")
            merger.append(file_path)
            # except Exception as e:
                # print(f"Error merging file {uploaded_file.file.name}: {e}")
        
        # Append VAT doc at the end (if not already merged)
        if vat_docs and vat_instance.file.path not in merger.pages:
            merger.append(vat_instance.file.path)

        # Save merged PDF
        buffer = BytesIO()
        merger.write(buffer)
        merger.close()
        
        merged_file = ContentFile(buffer.getvalue(), 'final_quotation.pdf')
        quotation.merged_pdf.save('quotation_merged.pdf', merged_file)

        # Set the content type and disposition headers for download
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=final_quotation.pdf'
        return response

    return render(request, 'pdfgenerator/upload.html', {
        'quotation': quotation,
        'excel_data': rows
    })



# def upload_pdfs(request, quotation_id):
#     quotation = get_object_or_404(generatequotation, id=quotation_id)
#     excel_path = os.path.join(settings.MEDIA_ROOT, quotation.excel.name)
#     data = pd.read_excel(excel_path, engine='openpyxl')
#     rows = data.iloc[:, 1].tolist()[:-3]  # Extract rows (excluding last 3)

#     if request.method == 'POST':
#         errors = []
        
#         # Handle file uploads
#         handle_file_uploads(request, quotation, rows, errors)

#         # Proceed to PDF merge after upload
#         return merge_uploaded_pdfs(quotation, errors)

#     return render(request, 'pdfgenerator/upload.html', {
#         'quotation': quotation,
#         'excel_data': rows
#     })


# # Function to handle file uploads
# def handle_file_uploads(request, quotation, rows, errors):
#     # Upload common files (Financial, BOQ, VAT)
#     file_mapping = {
#         'financial_docs': 'Financial Docs',
#         'boq_docs': 'BOQ',
#         'vat_docs': 'VAT'
#     }

#     for field, file_type in file_mapping.items():
#         uploaded_file = request.FILES.get(field)
#         if uploaded_file:
#             try:
#                 UploadedFile.objects.create(
#                     quotation=quotation,
#                     file=uploaded_file,
#                     file_type=file_type
#                 )
#             except Exception as e:
#                 errors.append(f"Failed to upload {file_type}: {str(e)}")

#     # Upload row-specific files (Catalogue, CE, ISO)
#     for row in rows:
#         for prefix, file_type in [('catalogue', 'Catalogue'), ('ce', 'CE'), ('iso', 'ISO')]:
#             uploaded_file = request.FILES.get(f'{prefix}_{row}')
#             if uploaded_file:
#                 try:
#                     UploadedFile.objects.create(
#                         quotation=quotation,
#                         file=uploaded_file,
#                         row_name=row,
#                         file_type=file_type
#                     )
#                 except Exception as e:
#                     errors.append(f"Failed to upload {file_type} for {row}: {str(e)}")


# # Function to merge PDFs
# def merge_uploaded_pdfs(quotation, errors):
#     uploaded_files = UploadedFile.objects.filter(quotation=quotation).order_by('file_type', 'uploaded_at')

#     merger = PdfMerger()
#     merged_files = set()

#     for uploaded_file in uploaded_files:
#         try:
#             file_path = uploaded_file.file.path
#             if uploaded_file.file_type == 'VAT' and file_path in merged_files:
#                 continue  # Prevent VAT duplication
#             merger.append(file_path)
#             merged_files.add(file_path)
#         except Exception as e:
#             errors.append(f"Failed to merge {uploaded_file.file_type}: {str(e)}")

#     buffer = BytesIO()
#     merger.write(buffer)
#     merger.close()

#     # Save merged PDF to the quotation
#     merged_file = ContentFile(buffer.getvalue(), 'final_quotation.pdf')
#     quotation.merged_pdf.save('quotation_merged.pdf', merged_file)

#     # Print errors if any
#     if errors:
#         for error in errors:
#             print(error)

#     return HttpResponse(buffer.getvalue(), content_type='application/pdf')
